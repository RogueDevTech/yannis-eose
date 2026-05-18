"""Generate Order Flow Test Plan as a Word (.docx) document."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc = Document()

# ── Styles ──────────────────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)

for level in range(1, 4):
    h = doc.styles[f"Heading {level}"]
    h.font.name = "Calibri"
    h.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)


def add_table(headers, rows):
    """Add a styled table to the document."""
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.style = "Light Grid Accent 1"
    # Header row
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.style = doc.styles["Normal"]
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    # Data rows
    for row_data in rows:
        row = t.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = str(val)
            for p in row.cells[i].paragraphs:
                p.style = doc.styles["Normal"]
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()


def add_checklist(items):
    """Add a checkbox checklist."""
    for item in items:
        p = doc.add_paragraph()
        p.style = doc.styles["List Bullet"]
        p.text = f"\u2610  {item}"
        p.paragraph_format.space_after = Pt(2)


# ── TITLE PAGE ──────────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading("Yannis EOSE", level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_heading("Order Flow Test Plan", level=1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
p = doc.add_paragraph("Manual QA test plan for the order submission pipeline.")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph("Split into 8 sections for team-based testing.")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
p = doc.add_paragraph("~120 unique test submissions  |  8 test sections  |  Ready-to-use mock data")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in p.runs:
    run.bold = True
doc.add_page_break()

# ── TABLE OF CONTENTS (manual) ─────────────────────────────────────────
doc.add_heading("Table of Contents", level=1)
toc_items = [
    "Pre-Test Setup",
    "Section A: Basic Order Count Accuracy (25 orders)",
    "Section B: Duplicate Detection (20 orders)",
    "Section C: Cross-Funnel Attribution (15 attempts)",
    "Section D: Cart Abandonment & Recovery (20 interactions)",
    "Section E: Rate Limiting & Edge Validation (15 attempts)",
    "Section F: Offline / Manual Order Entry (10 orders)",
    "Section G: Order Lifecycle State Transitions (10 orders)",
    "Section H: Payment Method Variants (5 orders)",
    "Phone Number Reference",
    "Assignment Sheet",
    "Post-Test Verification",
]
for i, item in enumerate(toc_items):
    doc.add_paragraph(f"{i+1}.  {item}")
doc.add_page_break()

# ── PRE-TEST SETUP ─────────────────────────────────────────────────────
doc.add_heading("Pre-Test Setup", level=1)

doc.add_heading("What You Need", level=2)
needs = [
    "Form URL: https://<edge-worker-domain>/form/<campaignId> — get the campaign ID from Admin > Forms",
    "At least 2 active campaigns tied to different Media Buyers (needed for cross-funnel tests in Section C)",
    "At least 2 products with stock in at least 1 logistics location",
    "Admin access to verify order counts, duplicate flags, and cart abandonments in the dashboard",
    "Logins for CS_CLOSER, HoCS, Logistics, and Finance roles (for Section G lifecycle tests)",
]
for n in needs:
    doc.add_paragraph(n, style="List Bullet")

doc.add_heading("Phone Number Rules", level=2)
doc.add_paragraph("All phones must be valid Nigerian format:")
doc.add_paragraph("08XXXXXXXXX — 11 digits, starts with 08, 07, or 09", style="List Bullet")
doc.add_paragraph("+2348XXXXXXXXX — 13 chars with country code", style="List Bullet")

doc.add_heading("How to Read the Test Data Tables", level=2)
doc.add_paragraph("Phone = paste into the form's phone field")
doc.add_paragraph("Name = paste into customer name field")
doc.add_paragraph("State = select from delivery state dropdown (if enabled on the form)")
doc.add_paragraph("Expected = what should happen after submission — verify this in the dashboard")

doc.add_heading("Tips for All Testers", level=2)
tips = [
    "Screenshot failures — if something doesn't match the expected result, screenshot the form error or dashboard state",
    "Note the exact time of each submission — helps correlate with temporal audit logs",
    "Don't clear browser data between tests in the same section — cookies/session matter for rate limiting",
    "Use a fresh incognito window when starting a different section to avoid cross-contamination",
]
for t in tips:
    doc.add_paragraph(t, style="List Bullet")
doc.add_page_break()

# ── SECTION A ───────────────────────────────────────────────────────────
doc.add_heading("Section A: Basic Order Count Accuracy", level=1)
p = doc.add_paragraph()
p.add_run("Assign to: ").bold = True
p.add_run("Someone with Admin dashboard access")
p = doc.add_paragraph()
p.add_run("Goal: ").bold = True
p.add_run("Submit 25 clean orders and verify the dashboard count matches exactly. No duplicates, no cart tricks, no cross-funnel — just straight submissions.")
p = doc.add_paragraph()
p.add_run("Estimated time: ").bold = True
p.add_run("30 minutes")

doc.add_heading("Before You Start", level=2)
doc.add_paragraph("1. Open the Admin dashboard and note the current total order count")
doc.add_paragraph("2. Note the current UNPROCESSED count")
doc.add_paragraph("3. Open the form URL in another tab")

doc.add_heading("Test Data — 25 Orders", level=2)

a_data = [
    ("A1", "Chinedu Okafor", "08031000001", "Lagos", "12 Admiralty Way, Lekki"),
    ("A2", "Amina Bello", "08031000002", "Kano", "5 Ibrahim Taiwo Rd"),
    ("A3", "Oluwaseun Adeyemi", "08031000003", "Oyo", "22 Ring Road, Ibadan"),
    ("A4", "Fatimah Abdullahi", "08031000004", "FCT (Abuja)", "8 Aso Drive, Maitama"),
    ("A5", "Emeka Nwosu", "08031000005", "Rivers", "15 Trans-Amadi Rd, PH"),
    ("A6", "Aisha Mohammed", "08031000006", "Kaduna", "3 Ahmadu Bello Way"),
    ("A7", "Tunde Bakare", "08031000007", "Lagos", "45 Allen Avenue, Ikeja"),
    ("A8", "Ngozi Eze", "08031000008", "Enugu", "10 Ogui Road"),
    ("A9", "Ibrahim Suleiman", "08031000009", "Kano", "7 Bompai Road"),
    ("A10", "Blessing Okonkwo", "08031000010", "Delta", "18 Warri-Sapele Rd"),
    ("A11", "Yusuf Garba", "08031000011", "Borno", "2 Shehu Laminu Way, Maiduguri"),
    ("A12", "Funke Adesanya", "08031000012", "Osun", "9 Station Road, Osogbo"),
    ("A13", "Obinna Chukwu", "08031000013", "Anambra", "6 New Market Rd, Onitsha"),
    ("A14", "Halima Usman", "08031000014", "Bauchi", "4 Jos Road, Bauchi"),
    ("A15", "Segun Ojo", "08031000015", "Ogun", "11 Abeokuta Expressway"),
    ("A16", "Chioma Igwe", "08031000016", "Imo", "20 Wetheral Road, Owerri"),
    ("A17", "Musa Danjuma", "08031000017", "Plateau", "14 Bukuru Road, Jos"),
    ("A18", "Adaeze Obi", "08031000018", "Edo", "7 Sapele Road, Benin City"),
    ("A19", "Kabiru Aliyu", "08031000019", "Sokoto", "1 Sultan Ibrahim Rd"),
    ("A20", "Kemi Adekunle", "08031000020", "Ekiti", "5 Fajuyi Road, Ado-Ekiti"),
    ("A21", "Uche Nnamdi", "08031000021", "Abia", "33 Aba-Owerri Road"),
    ("A22", "Zainab Bala", "08031000022", "Niger", "8 Paiko Road, Minna"),
    ("A23", "Dare Ogunlade", "08031000023", "Ondo", "16 Oyemekun Road, Akure"),
    ("A24", "Ifeoma Aneke", "08031000024", "Cross River", "12 Marian Road, Calabar"),
    ("A25", "Sani Abubakar", "08031000025", "Gombe", "3 Bauchi Road, Gombe"),
]
add_table(
    ["#", "Customer Name", "Phone", "Delivery State", "Address"],
    a_data,
)
doc.add_paragraph("Expected result for ALL 25: Status = UNPROCESSED, orderSource = edge-form")

doc.add_heading("Verification Checklist", level=2)
add_checklist([
    "Total order count increased by exactly 25",
    "UNPROCESSED count increased by exactly 25",
    "Each order has a unique Order ID (sorted by submission time)",
    "All 25 orders visible in Admin > Orders list",
    "Phone numbers are masked (not visible in raw form) for non-admin roles",
    "Each order has the correct product and offer attached",
    "orderSource = edge-form on all 25",
])
doc.add_page_break()

# ── SECTION B ───────────────────────────────────────────────────────────
doc.add_heading("Section B: Duplicate Detection", level=1)
p = doc.add_paragraph()
p.add_run("Assign to: ").bold = True
p.add_run("Someone who can submit orders AND check the CS/HoCS dashboard")
p = doc.add_paragraph()
p.add_run("Goal: ").bold = True
p.add_run("Trigger all duplicate scenarios and verify correct flagging.")
p = doc.add_paragraph()
p.add_run("Estimated time: ").bold = True
p.add_run("45 minutes")

doc.add_heading("Step 1 — Submit Original Orders (no flags expected)", level=2)
b1_data = [
    ("B1", "Adamu Kolo", "08032000001", "Original order", "No flag"),
    ("B2", "Grace Emenike", "08032000002", "Original order", "No flag"),
    ("B3", "Rasheed Lawal", "08032000003", "Original order", "No flag"),
    ("B4", "Stella Okoye", "08032000004", "Original order", "No flag"),
    ("B5", "John Akpan", "08032000005", "Original order", "No flag"),
]
add_table(["#", "Customer Name", "Phone", "Notes", "Expected Flag"], b1_data)

doc.add_heading("Step 2 — Submit Duplicates (same phone, same product, within minutes)", level=2)
doc.add_paragraph("Use the SAME form (same campaign, same product) and the SAME phone numbers as Step 1.")
b2_data = [
    ("B6", "Adamu Kolo", "08032000001", "Same phone as B1", "FLAGGED"),
    ("B7", "Grace E.", "08032000002", "Same phone as B2", "FLAGGED"),
    ("B8", "R. Lawal", "08032000003", "Same phone as B3", "FLAGGED"),
    ("B9", "Stella O.", "08032000004", "Same phone as B4", "FLAGGED"),
    ("B10", "J. Akpan", "08032000005", "Same phone as B5", "FLAGGED"),
]
add_table(["#", "Customer Name", "Phone", "Notes", "Expected Flag"], b2_data)

doc.add_heading("Step 3 — Submit Same Phones with a DIFFERENT Product", level=2)
doc.add_paragraph("Switch to a different product on the same form (or a different campaign with a different product). Use the same phone numbers again.")
b3_data = [
    ("B11", "Adamu Kolo", "08032000001", "Different product", "FLAGGED (same phone within 24h)"),
    ("B12", "Grace Emenike", "08032000002", "Different product", "FLAGGED"),
    ("B13", "Rasheed Lawal", "08032000003", "Different product", "FLAGGED"),
    ("B14", "Stella Okoye", "08032000004", "Different product", "FLAGGED"),
    ("B15", "John Akpan", "08032000005", "Different product", "FLAGGED"),
]
add_table(["#", "Customer Name", "Phone", "Notes", "Expected Flag"], b3_data)

doc.add_heading("Step 4 — Test Duplicate Management Actions (HoCS)", level=2)
doc.add_paragraph("Log in as HoCS and manage the flagged orders:")
add_table(
    ["Action", "Order", "Expected Result"],
    [
        ("DISMISS", "B6", "isDuplicate changes to DISMISSED, order continues normal flow"),
        ("MERGE", "B7 into B2", "B7 becomes MERGED, B2's item quantities combined"),
        ("Leave alone", "B8", "Stays FLAGGED — visible in HoCS duplicate queue"),
    ],
)

doc.add_heading("Verification Checklist", level=2)
add_checklist([
    "B6-B10 all show FLAGGED badge in order detail",
    "duplicateOfId links to the correct original order (B1-B5)",
    "HoCS can see flagged orders in duplicate filter",
    "DISMISS clears the flag successfully",
    "MERGE combines quantities and marks as MERGED",
    "Total order count includes ALL orders (flagged + unflagged)",
    "B11-B15 also flagged (same phone within 24h regardless of product)",
])

doc.add_heading("Note on POSSIBLY_DUPLICATE (soft flag)", level=2)
doc.add_paragraph(
    "To test the POSSIBLY_DUPLICATE flag (24h-30d window), you would need to wait 24+ hours "
    "after B1-B5 and resubmit with the same phones. For practical testing, verifying the FLAGGED "
    "behavior (B6-B15) is sufficient. The soft flag uses the same code path with a different time window."
)
doc.add_page_break()

# ── SECTION C ───────────────────────────────────────────────────────────
doc.add_heading("Section C: Cross-Funnel Attribution", level=1)
p = doc.add_paragraph()
p.add_run("Assign to: ").bold = True
p.add_run("Someone with access to 2 different campaign form URLs owned by different Media Buyers")
p = doc.add_paragraph()
p.add_run("Goal: ").bold = True
p.add_run("Verify that the first Media Buyer keeps attribution and the second MB's attempt is recorded but NO order is created.")
p = doc.add_paragraph()
p.add_run("Estimated time: ").bold = True
p.add_run("30 minutes")

doc.add_heading("Setup Required", level=2)
doc.add_paragraph("Campaign X = owned by Media Buyer Alpha (get form URL from Admin > Forms)")
doc.add_paragraph("Campaign Y = owned by Media Buyer Beta (different MB), must have the SAME product as Campaign X")
doc.add_paragraph("Both campaigns must be ACTIVE with at least 1 shared product.")

doc.add_heading("Step 1 — Submit on Campaign X (MB-Alpha wins attribution)", level=2)
c1_data = [
    ("C1", "Taiwo Adeniyi", "08033000001", "Campaign X", "Order created, MB-Alpha gets it"),
    ("C2", "Nkechi Uzoma", "08033000002", "Campaign X", "Order created, MB-Alpha gets it"),
    ("C3", "Hassan Yusuf", "08033000003", "Campaign X", "Order created, MB-Alpha gets it"),
    ("C4", "Vivian Edet", "08033000004", "Campaign X", "Order created, MB-Alpha gets it"),
    ("C5", "Olu Bankole", "08033000005", "Campaign X", "Order created, MB-Alpha gets it"),
]
add_table(["#", "Customer Name", "Phone", "Submit On", "Expected"], c1_data)

doc.add_heading("Step 2 — Immediately resubmit SAME phones on Campaign Y (MB-Beta loses)", level=2)
doc.add_paragraph("Submit the exact same phone numbers on Campaign Y. These should NOT create orders.")
c2_data = [
    ("C6", "Taiwo Adeniyi", "08033000001", "Campaign Y", "NO order. Cross-funnel attempt recorded."),
    ("C7", "Nkechi Uzoma", "08033000002", "Campaign Y", "NO order. Cross-funnel attempt recorded."),
    ("C8", "Hassan Yusuf", "08033000003", "Campaign Y", "NO order. Cross-funnel attempt recorded."),
    ("C9", "Vivian Edet", "08033000004", "Campaign Y", "NO order. Cross-funnel attempt recorded."),
    ("C10", "Olu Bankole", "08033000005", "Campaign Y", "NO order. Cross-funnel attempt recorded."),
]
add_table(["#", "Customer Name", "Phone", "Submit On", "Expected"], c2_data)

doc.add_heading("Step 3 — Submit NEW phones on Campaign Y (MB-Beta gets these)", level=2)
c3_data = [
    ("C11", "Folake Ige", "08033000011", "Campaign Y", "Order created, MB-Beta gets it"),
    ("C12", "Murtala Shehu", "08033000012", "Campaign Y", "Order created, MB-Beta gets it"),
    ("C13", "Ebere Udeh", "08033000013", "Campaign Y", "Order created, MB-Beta gets it"),
    ("C14", "Garba Aliyu", "08033000014", "Campaign Y", "Order created, MB-Beta gets it"),
    ("C15", "Ada Nwachukwu", "08033000015", "Campaign Y", "Order created, MB-Beta gets it"),
]
add_table(["#", "Customer Name", "Phone", "Submit On", "Expected"], c3_data)

doc.add_heading("Verification Checklist", level=2)
add_checklist([
    "Only 10 orders created total (C1-C5 + C11-C15), NOT 15",
    "C6-C10 appear in cross_funnel_attempts (Admin > Marketing analytics or DB)",
    "MB-Alpha's order count = 5 (C1-C5)",
    "MB-Beta's order count = 5 (C11-C15), NOT 10",
    "MB-Beta can see their cross-funnel attempt count in their dashboard",
    "Cross-funnel attempts are NOT counted in total order count",
    "originalOrderId on each attempt points to the correct C1-C5 order",
    "originalMediaBuyerId = MB-Alpha's ID on all 5 attempts",
])
doc.add_page_break()

# ── SECTION D ───────────────────────────────────────────────────────────
doc.add_heading("Section D: Cart Abandonment & Recovery", level=1)
p = doc.add_paragraph()
p.add_run("Assign to: ").bold = True
p.add_run("Someone to partially fill forms + someone with CS dashboard access to verify")
p = doc.add_paragraph()
p.add_run("Goal: ").bold = True
p.add_run("Verify carts are captured on partial form fill and correctly transition to ABANDONED or CONVERTED.")
p = doc.add_paragraph()
p.add_run("Estimated time: ").bold = True
p.add_run("30 minutes + 10 minute wait for the abandonment cron")

doc.add_heading("How Cart Saving Works", level=2)
doc.add_paragraph(
    "The form auto-saves to the cart when the user has filled in name + phone but hasn't submitted. "
    "The save triggers on field blur/change events. After 5 minutes of inactivity, a background cron "
    "marks PENDING carts as ABANDONED."
)

doc.add_heading("Step 1 — Create Abandoned Carts (fill partially, then close the tab)", level=2)
d1_data = [
    ("D1", "Patience Obiora", "08034000001", "Name + Phone only", "Close tab"),
    ("D2", "Ikenna Okafor", "08034000002", "Name + Phone + Address", "Close tab"),
    ("D3", "Mariam Isah", "08034000003", "Name + Phone + State (Lagos)", "Close tab"),
    ("D4", "Chibuzor Agu", "08034000004", "Name + Phone + All fields", "Close tab"),
    ("D5", "Yetunde Balogun", "08034000005", "Name + Phone only", "Close tab"),
    ("D6", "Abdulrahman Ismail", "08034000006", "Name + Phone + Email", "Close tab"),
    ("D7", "Toyin Fasasi", "08034000007", "Name + Phone + Notes", "Close tab"),
    ("D8", "Nnamdi Okoro", "08034000008", "Name + Phone + Gender", "Close tab"),
    ("D9", "Safiya Gambo", "08034000009", "Name + Phone only", "Close tab"),
    ("D10", "Bisi Adeleke", "08034000010", "Name + Phone + Address + State", "Close tab"),
]
add_table(["#", "Customer Name", "Phone", "Fields to Fill", "Then..."], d1_data)
doc.add_paragraph("Expected: All 10 carts appear as PENDING, then transition to ABANDONED after ~5 minutes.")

doc.add_heading("Step 2 — Recover Some Carts (go back and submit)", level=2)
doc.add_paragraph("Wait at least 5 minutes for D1-D5 to become ABANDONED, then reopen the form and submit with the same phone:")
d2_data = [
    ("D11", "Patience Obiora", "08034000001", "Complete and submit", "Cart -> CONVERTED, order created with cartId"),
    ("D12", "Ikenna Okafor", "08034000002", "Complete and submit", "Cart -> CONVERTED, order created with cartId"),
    ("D13", "Mariam Isah", "08034000003", "Complete and submit", "Cart -> CONVERTED, order created with cartId"),
]
add_table(["#", "Customer Name", "Phone", "Action", "Expected"], d2_data)

doc.add_heading("Step 3 — Leave the Rest as Permanently Abandoned", level=2)
doc.add_paragraph("Do NOT go back for these. They should stay ABANDONED forever:")
d3_data = [
    ("D14", "08034000004", "Chibuzor Agu", "Stays ABANDONED"),
    ("D15", "08034000005", "Yetunde Balogun", "Stays ABANDONED"),
    ("D16", "08034000006", "Abdulrahman Ismail", "Stays ABANDONED"),
    ("D17", "08034000007", "Toyin Fasasi", "Stays ABANDONED"),
    ("D18", "08034000008", "Nnamdi Okoro", "Stays ABANDONED"),
    ("D19", "08034000009", "Safiya Gambo", "Stays ABANDONED"),
    ("D20", "08034000010", "Bisi Adeleke", "Stays ABANDONED"),
]
add_table(["#", "Phone", "Name", "Expected"], d3_data)

doc.add_heading("Verification Checklist", level=2)
add_checklist([
    "All 10 carts appear in CS cart abandonment view within 5-10 minutes",
    "Carts show progressive field data (address, state, etc. where filled)",
    "D1-D10 transition from PENDING to ABANDONED after the 5-min cron",
    "D11-D13 transition from ABANDONED to CONVERTED after form submission",
    "Converted orders have cartId populated (visible in order detail)",
    "Abandonment rate = 7/10 (70%) matches dashboard stats",
    "Recovery rate = 3/10 (30%) matches dashboard stats",
    "Total order count only includes the 3 converted, not the 7 abandoned",
])
doc.add_page_break()

# ── SECTION E ───────────────────────────────────────────────────────────
doc.add_heading("Section E: Rate Limiting & Edge Validation", level=1)
p = doc.add_paragraph()
p.add_run("Assign to: ").bold = True
p.add_run("A single person submitting rapidly from one browser/IP")
p = doc.add_paragraph()
p.add_run("Goal: ").bold = True
p.add_run("Verify the edge worker blocks abuse without blocking legitimate orders.")
p = doc.add_paragraph()
p.add_run("Estimated time: ").bold = True
p.add_run("20 minutes")

doc.add_heading("Test 1 — Rate Limit (5 submissions per 5 minutes per IP)", level=2)
doc.add_paragraph("Submit these as fast as possible from the same browser:")
e1_data = [
    ("E1", "Test Rate A", "08035000001", "Immediately", "Success"),
    ("E2", "Test Rate B", "08035000002", "+10 seconds", "Success"),
    ("E3", "Test Rate C", "08035000003", "+20 seconds", "Success (CAPTCHA may appear)"),
    ("E4", "Test Rate D", "08035000004", "+30 seconds", "Success (CAPTCHA required)"),
    ("E5", "Test Rate E", "08035000005", "+40 seconds", "Success (CAPTCHA required)"),
    ("E6", "Test Rate F", "08035000006", "+50 seconds", "BLOCKED — rate limit exceeded"),
]
add_table(["#", "Name", "Phone", "Timing", "Expected"], e1_data)

doc.add_heading("Test 2 — Edge Dedup (same phone blocked for 6 hours)", level=2)
doc.add_paragraph("Wait for rate limit to reset (5 minutes), then test dedup:")
e2_data = [
    ("E7", "Dedup Test A", "08035000011", "First submission", "Success"),
    ("E8", "Dedup Test A", "08035000011", "Same phone, immediate retry", "BLOCKED — 6h dedup"),
    ("E9", "Dedup Test B", "08035000012", "New phone", "Success"),
    ("E10", "Dedup Test B", "08035000012", "Same phone, 1 min later", "BLOCKED — 6h dedup"),
]
add_table(["#", "Name", "Phone", "Notes", "Expected"], e2_data)

doc.add_heading("Test 3 — Validation Failures", level=2)
e3_data = [
    ("E11", "A", "08035000021", "Name too short (min 2 chars)", "Validation error"),
    ("E12", "Valid Name", "12345", "Invalid phone format", "Validation error"),
    ("E13", "Valid Name", "08035000023", "No product selected (empty items)", "Validation error"),
    ("E14", "Valid Name", "+1234567890", "Non-Nigerian phone number", "Validation error"),
    ("E15", "Valid Name", "08035000025", "Everything valid", "Success"),
]
add_table(["#", "Name", "Phone", "Problem", "Expected"], e3_data)

doc.add_heading("Verification Checklist", level=2)
add_checklist([
    "First 5 submissions succeed (E1-E5)",
    "CAPTCHA appears after 3rd submission",
    "6th submission blocked with rate limit error message",
    "Edge dedup blocks same phone resubmission within 6 hours (E8, E10)",
    "Validation errors show user-friendly messages (E11-E14)",
    "After 5 minutes, rate limit resets and you can submit again",
    "None of the blocked/failed submissions appear in the orders table",
])
doc.add_page_break()

# ── SECTION F ───────────────────────────────────────────────────────────
doc.add_heading("Section F: Offline / Manual Order Entry", level=1)
p = doc.add_paragraph()
p.add_run("Assign to: ").bold = True
p.add_run("Someone with CS_CLOSER or HoCS role login")
p = doc.add_paragraph()
p.add_run("Goal: ").bold = True
p.add_run("Verify offline/manual order creation works and is correctly tagged as orderSource = offline.")
p = doc.add_paragraph()
p.add_run("Estimated time: ").bold = True
p.add_run("20 minutes")

doc.add_heading("How Manual Entry Works", level=2)
doc.add_paragraph(
    "CS agents can create orders manually from the CS dashboard (e.g., for walk-in customers, "
    "phone inquiries, social media leads). The CS enters the raw phone number and the API hashes "
    "it server-side. These orders are tagged as orderSource = offline."
)

doc.add_heading("Test Data — 10 Manual Orders", level=2)
f_data = [
    ("F1", "Chukwudi Nnaji", "08036000001", "Lagos", "Walk-in customer"),
    ("F2", "Hajara Musa", "08036000002", "Kano", "Phone call inquiry"),
    ("F3", "Solomon Udo", "08036000003", "Rivers", "WhatsApp lead"),
    ("F4", "Rukayat Adeyinka", "08036000004", "Oyo", "Referral"),
    ("F5", "Patrick Eke", "08036000005", "FCT (Abuja)", "Instagram DM lead"),
    ("F6", "Bilkisu Tanko", "08036000006", "Kaduna", "Repeat customer"),
    ("F7", "Chidera Anyanwu", "08036000007", "Enugu", "Phone call"),
    ("F8", "Abubakar Sadiq", "08036000008", "Bauchi", "Walk-in"),
    ("F9", "Nneka Okorie", "08036000009", "Anambra", "Social media lead"),
    ("F10", "Mustapha Bello", "08036000010", "Lagos", "Referral from F1"),
]
add_table(["#", "Customer Name", "Phone", "Delivery State", "Source Notes"], f_data)

doc.add_heading("Verification Checklist", level=2)
add_checklist([
    "All 10 orders created successfully from CS dashboard",
    "orderSource = offline on all 10 (not edge-form)",
    "Phone is hashed server-side (CS enters raw phone, API hashes it)",
    "Orders appear in the same order list as edge-form orders",
    "Duplicate detection still works (try submitting F1's phone again — should flag)",
    "Offline orders go through normal state machine (UNPROCESSED onward)",
])
doc.add_page_break()

# ── SECTION G ───────────────────────────────────────────────────────────
doc.add_heading("Section G: Order Lifecycle State Transitions", level=1)
p = doc.add_paragraph()
p.add_run("Assign to: ").bold = True
p.add_run("Multiple people coordinating — CS handles early states, Logistics handles dispatch, Finance handles remittance")
p = doc.add_paragraph()
p.add_run("Goal: ").bold = True
p.add_run("Walk 10 orders through different lifecycle paths and verify every transition gate.")
p = doc.add_paragraph()
p.add_run("Estimated time: ").bold = True
p.add_run("60 minutes")

doc.add_heading("Use orders from Section A (or create fresh ones). Each tests a different path:", level=2)

g_data = [
    ("Order 1", "UNPROCESSED -> CS_ASSIGNED -> CS_ENGAGED -> CONFIRMED -> AGENT_ASSIGNED -> DISPATCHED -> IN_TRANSIT -> DELIVERED -> REMITTED", "Full happy path — all gates"),
    ("Order 2", "UNPROCESSED -> CANCELLED", "Early cancel (reason required, 10+ chars)"),
    ("Order 3", "CS_ASSIGNED -> CANCELLED", "Cancel after assignment"),
    ("Order 4", "CS_ENGAGED -> CANCELLED", "Cancel after engagement"),
    ("Order 5", "... -> IN_TRANSIT -> RETURNED -> RESTOCKED", "Return + restock path"),
    ("Order 6", "... -> IN_TRANSIT -> RETURNED -> WRITTEN_OFF", "Return + write-off (damage note required)"),
    ("Order 7", "... -> IN_TRANSIT -> PARTIALLY_DELIVERED -> REMITTED", "Partial delivery (specify qty)"),
    ("Order 8", "CONFIRMED -> AGENT_ASSIGNED -> DELIVERED (CS proxy)", "CS proxy delivery (note 10+ chars required)"),
    ("Order 9", "AGENT_ASSIGNED -> AGENT_ASSIGNED (different location)", "Re-assignment (new location needs stock)"),
    ("Order 10", "Full path — verify all timestamp fields", "Check confirmedAt, allocatedAt, dispatchedAt, deliveredAt"),
]
add_table(["Order", "Path to Test", "Key Gate"], g_data)

doc.add_heading("State Transition Gates to Verify", level=2)
add_table(
    ["Transition", "Gate / Requirement"],
    [
        ("Any -> CANCELLED", "Requires reason note, minimum 10 characters"),
        ("CS_ENGAGED -> CONFIRMED", "Requires qualifying call (VOIP 15s+ or manual call log)"),
        ("CONFIRMED -> AGENT_ASSIGNED", "3PL location must have available stock"),
        ("AGENT_ASSIGNED -> DISPATCHED", "Rider must be assigned"),
        ("IN_TRANSIT -> DELIVERED", "OTP match required (or SuperAdmin override)"),
        ("IN_TRANSIT -> RETURNED", "Mandatory return reason"),
        ("RETURNED -> WRITTEN_OFF", "Mandatory damage note"),
        ("DELIVERED -> REMITTED", "Accountant/Finance role ONLY — CS can NEVER do this"),
        ("AGENT_ASSIGNED -> DELIVERED (CS proxy)", "Mandatory delivery note, 10+ chars"),
    ],
)

doc.add_heading("Negative Tests (these should all FAIL)", level=2)
add_checklist([
    "Try UNPROCESSED -> CONFIRMED directly (should fail — can't skip states)",
    "Try UNPROCESSED -> DELIVERED directly (should fail)",
    "Try CS_ENGAGED -> CONFIRMED without a call log (should fail)",
    "Try CONFIRMED -> AGENT_ASSIGNED at a location with 0 stock (should fail)",
    "Try CS role marking REMITTED (should fail — accountant only)",
    "Try CANCELLED -> any other state (should fail — terminal state)",
])

doc.add_heading("Verification Checklist", level=2)
add_checklist([
    "No state can be skipped",
    "All transition gates enforce their requirements",
    "Every transition logged in temporal audit trail (check with Admin)",
    "Timestamp fields populated at correct transitions",
    "Stock reserved at CONFIRMED, allocated at AGENT_ASSIGNED, deducted at DELIVERED",
    "REMITTED only accessible to Finance/Accountant role",
    "All negative tests fail with appropriate error messages",
])
doc.add_page_break()

# ── SECTION H ───────────────────────────────────────────────────────────
doc.add_heading("Section H: Payment Method Variants", level=1)
p = doc.add_paragraph()
p.add_run("Assign to: ").bold = True
p.add_run("Someone testing the form with payment options enabled")
p = doc.add_paragraph()
p.add_run("Goal: ").bold = True
p.add_run("Verify both Cash on Delivery and Pay Online paths work correctly.")
p = doc.add_paragraph()
p.add_run("Estimated time: ").bold = True
p.add_run("15 minutes")

doc.add_heading("Test Data", level=2)
h_data = [
    ("H1", "Emeka Uche", "08037000001", "PAY_ON_DELIVERY", "(not needed)", "Normal order flow"),
    ("H2", "Amaka Peters", "08037000002", "PAY_ON_DELIVERY", "amaka@test.com", "Normal flow (email stored but optional)"),
    ("H3", "Ifeanyi Obi", "08037000003", "PAY_ONLINE", "ifeanyi@test.com", "Redirects to Paystack payment"),
    ("H4", "Ngozi Amadi", "08037000004", "PAY_ONLINE", "(leave empty)", "VALIDATION ERROR — email required"),
    ("H5", "Bola Adewale", "08037000005", "PAY_ONLINE", "bola@test.com", "Redirects to Paystack payment"),
]
add_table(["#", "Name", "Phone", "Payment Method", "Email", "Expected"], h_data)

doc.add_heading("Verification Checklist", level=2)
add_checklist([
    "PAY_ON_DELIVERY orders have paymentStatus = PENDING",
    "PAY_ONLINE orders redirect to Paystack payment page",
    "PAY_ONLINE without email is rejected with validation error",
    "paymentMethod stored correctly on order record",
    "paymentReference populated after Paystack callback (for online payments)",
])
doc.add_page_break()

# ── PHONE NUMBER REFERENCE ──────────────────────────────────────────────
doc.add_heading("Phone Number Reference — All Sections", level=1)
doc.add_paragraph("Use this to avoid accidental phone collisions between sections:")
add_table(
    ["Section", "Phone Range", "Count"],
    [
        ("A — Basic Count", "08031000001 – 08031000025", "25"),
        ("B — Duplicates", "08032000001 – 08032000005", "5 unique (20 submissions)"),
        ("C — Cross-Funnel", "08033000001 – 08033000015", "15"),
        ("D — Cart Abandon", "08034000001 – 08034000010", "10"),
        ("E — Rate Limit", "08035000001 – 08035000025", "15"),
        ("F — Offline", "08036000001 – 08036000010", "10"),
        ("G — Lifecycle", "(reuses Section A orders)", "0 new"),
        ("H — Payment", "08037000001 – 08037000005", "5"),
        ("TOTAL", "", "~120 unique submissions"),
    ],
)
doc.add_page_break()

# ── ASSIGNMENT SHEET ────────────────────────────────────────────────────
doc.add_heading("Assignment Sheet", level=1)
doc.add_paragraph("Fill in tester names and share the relevant sections with each person:")
add_table(
    ["Section", "Description", "Tester Name", "Est. Time", "Roles Needed"],
    [
        ("A", "Basic order count (25 orders)", "", "30 min", "Form URL + Admin dashboard"),
        ("B", "Duplicate detection (20 orders)", "", "45 min", "Form URL + HoCS dashboard"),
        ("C", "Cross-funnel attribution (15)", "", "30 min", "2 campaign URLs (different MBs)"),
        ("D", "Cart abandonment (20)", "", "40 min", "Form URL + CS dashboard"),
        ("E", "Rate limiting & validation (15)", "", "20 min", "Form URL (single browser)"),
        ("F", "Offline/manual entry (10)", "", "20 min", "CS_CLOSER or HoCS login"),
        ("G", "Lifecycle transitions (10)", "", "60 min", "Admin + CS + Logistics + Finance"),
        ("H", "Payment variants (5)", "", "15 min", "Form URL with payment enabled"),
    ],
)
doc.add_page_break()

# ── POST-TEST VERIFICATION ──────────────────────────────────────────────
doc.add_heading("Post-Test Verification (Admin)", level=1)
doc.add_paragraph("After all sections are complete, verify these totals in the Admin dashboard:")
add_table(
    ["Check", "Expected Value"],
    [
        ("Total orders created", "~83 (25 A + 15 B + 10 C + 3 D-converted + ~4 E + 10 F + ~4 H)"),
        ("Orders with isDuplicate = FLAGGED", "10 (B6-B15)"),
        ("Orders with orderSource = edge-form", "All except Section F"),
        ("Orders with orderSource = offline", "10 (Section F only)"),
        ("Cross-funnel attempts recorded", "5 (C6-C10) — NOT in orders table"),
        ("Cart abandonments (ABANDONED)", "7 (D14-D20)"),
        ("Cart abandonments (CONVERTED)", "3 (D11-D13)"),
        ("Orders with cartId populated", "3 (D11-D13)"),
        ("Rate-limited submissions", "NOT in orders table"),
        ("Validation failures", "NOT in orders table"),
    ],
)

doc.add_heading("Final Sign-Off", level=2)
doc.add_paragraph()
add_table(
    ["Section", "Pass / Fail", "Notes", "Tested By", "Date"],
    [
        ("A — Order Count", "", "", "", ""),
        ("B — Duplicates", "", "", "", ""),
        ("C — Cross-Funnel", "", "", "", ""),
        ("D — Cart Abandonment", "", "", "", ""),
        ("E — Rate Limiting", "", "", "", ""),
        ("F — Offline Entry", "", "", "", ""),
        ("G — Lifecycle", "", "", "", ""),
        ("H — Payments", "", "", "", ""),
    ],
)

# ── SAVE ────────────────────────────────────────────────────────────────
output_path = os.path.join(os.path.dirname(__file__), "Order-Flow-Test-Plan.docx")
doc.save(output_path)
print(f"Document saved to: {output_path}")
